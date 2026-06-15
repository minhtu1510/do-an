"""
Risk Assessment Report Generator

Generate professional pentest reports in PDF and DOCX formats
"""

import os
import json
from datetime import datetime
from typing import Optional
from pathlib import Path

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, Image, KeepTogether
    )
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ...domain.risk_assessment import RiskAssessmentReport, RiskLevel


class ReportGenerator:
    """Generate risk assessment reports in various formats"""

    def __init__(self, output_dir: str = "/home/user/S7.Pwn/reports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_pdf_report(
        self,
        report: RiskAssessmentReport,
        filename: Optional[str] = None
    ) -> str:
        """
        Generate PDF pentest report

        Returns: Path to generated PDF file
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab is required for PDF generation. Install with: pip install reportlab")

        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"risk_assessment_{timestamp}.pdf"

        filepath = self.output_dir / filename

        # Create PDF document
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=A4,
            topMargin=0.75*inch,
            bottomMargin=0.75*inch,
            leftMargin=0.75*inch,
            rightMargin=0.75*inch
        )

        # Container for flowable elements
        story = []
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )

        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#2c3e50'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )

        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#34495e'),
            spaceAfter=10,
            spaceBefore=10,
            fontName='Helvetica-Bold'
        )

        # Title Page
        story.append(Spacer(1, 1.5*inch))
        story.append(Paragraph("OT/ICS SECURITY", title_style))
        story.append(Paragraph("PENETRATION TEST REPORT", title_style))
        story.append(Spacer(1, 0.5*inch))

        # Risk level badge
        risk_color = self._get_risk_color(report.overall_risk_level)
        risk_text = f'<para align="center" fontSize="18" textColor="{risk_color}"><b>Risk Level: {report.overall_risk_level.value}</b></para>'
        story.append(Paragraph(risk_text, styles['Normal']))
        story.append(Spacer(1, 0.3*inch))

        # Report info
        info_data = [
            ["Report ID:", report.report_id],
            ["Assessment Date:", report.scan_timestamp.strftime("%Y-%m-%d %H:%M:%S")],
            ["Total Devices Scanned:", str(report.total_devices)],
            ["Overall Risk Score:", f"{report.overall_risk_score:.2f}/100"],
        ]
        info_table = Table(info_data, colWidths=[2.5*inch, 3.5*inch])
        info_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, -1), 'Helvetica', 10),
            ('FONT', (0, 0), (0, -1), 'Helvetica-Bold', 10),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(info_table)

        story.append(PageBreak())

        # Executive Summary
        story.append(Paragraph("1. EXECUTIVE SUMMARY", heading1_style))
        story.append(Spacer(1, 0.2*inch))

        exec_summary = f"""
        This report presents the results of a comprehensive security assessment of the OT/ICS network infrastructure.
        The assessment identified <b>{report.total_devices} devices</b>, with an overall risk rating of
        <b><font color="{risk_color}">{report.overall_risk_level.value}</font></b> (score: {report.overall_risk_score:.2f}/100).
        """
        story.append(Paragraph(exec_summary, styles['BodyText']))
        story.append(Spacer(1, 0.15*inch))

        # Key Findings Summary
        story.append(Paragraph("Key Findings:", heading2_style))
        key_findings_data = [
            ['Severity', 'Count', 'Status'],
            ['Critical', str(report.critical_count), '🔴 Immediate Action Required'],
            ['High', str(report.high_count), '🟠 Urgent'],
            ['Medium', str(len([f for f in report.all_findings if f.severity.value == "MEDIUM"])), '🟡 Important'],
            ['Low', str(len([f for f in report.all_findings if f.severity.value == "LOW"])), '🟢 Monitor'],
        ]
        findings_table = Table(key_findings_data, colWidths=[1.5*inch, 1*inch, 3.5*inch])
        findings_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')])
        ]))
        story.append(findings_table)

        story.append(PageBreak())

        # Assessment Details
        story.append(Paragraph("2. DETAILED ASSESSMENT", heading1_style))

        # Network Security
        story.append(Paragraph("2.1 Network Security Assessment", heading2_style))
        net_score = report.network_assessment.score
        net_risk = RiskLevel.from_score(100 - net_score)
        story.append(Paragraph(
            f"<b>Score:</b> {net_score:.1f}/100 | <b>Risk Level:</b> <font color=\"{self._get_risk_color(net_risk)}\">{net_risk.value}</font>",
            styles['BodyText']
        ))
        story.append(Spacer(1, 0.1*inch))

        if report.network_assessment.findings:
            story.append(Paragraph("<b>Findings:</b>", styles['BodyText']))
            for finding in report.network_assessment.findings[:5]:
                story.append(Paragraph(
                    f"• {finding.title}: {finding.description}",
                    styles['BodyText']
                ))
        story.append(Spacer(1, 0.15*inch))

        # Device Security
        story.append(Paragraph("2.2 Device Security Assessment", heading2_style))
        dev_score = report.device_assessment.score
        dev_risk = RiskLevel.from_score(100 - dev_score)
        story.append(Paragraph(
            f"<b>Score:</b> {dev_score:.1f}/100 | <b>Risk Level:</b> <font color=\"{self._get_risk_color(dev_risk)}\">{dev_risk.value}</font>",
            styles['BodyText']
        ))
        story.append(Spacer(1, 0.1*inch))

        if report.device_assessment.findings:
            story.append(Paragraph("<b>Findings:</b>", styles['BodyText']))
            for finding in report.device_assessment.findings[:5]:
                story.append(Paragraph(
                    f"• {finding.title}",
                    styles['BodyText']
                ))
        story.append(Spacer(1, 0.15*inch))

        # Vulnerability Assessment
        story.append(Paragraph("2.3 Vulnerability Assessment", heading2_style))
        vuln_score = report.vulnerability_assessment.score
        vuln_risk = RiskLevel.from_score(100 - vuln_score)
        story.append(Paragraph(
            f"<b>Score:</b> {vuln_score:.1f}/100 | <b>Risk Level:</b> <font color=\"{self._get_risk_color(vuln_risk)}\">{vuln_risk.value}</font>",
            styles['BodyText']
        ))
        story.append(Spacer(1, 0.1*inch))

        if report.critical_findings:
            story.append(Paragraph("<b>Critical Vulnerabilities:</b>", styles['BodyText']))
            for finding in report.critical_findings[:5]:
                cve_text = f" ({finding.cve_id})" if finding.cve_id else ""
                story.append(Paragraph(
                    f"• <b>{finding.title}</b>{cve_text}: {finding.description}",
                    styles['BodyText']
                ))

        story.append(PageBreak())

        # Compliance Assessment
        story.append(Paragraph("2.4 Compliance Assessment", heading2_style))
        comp_score = report.compliance_assessment.score
        story.append(Paragraph(
            f"<b>Compliance Score:</b> {comp_score:.1f}%",
            styles['BodyText']
        ))
        story.append(Spacer(1, 0.1*inch))

        if report.compliance_status:
            for comp_status in report.compliance_status:
                story.append(Paragraph(
                    f"<b>{comp_status.framework}:</b> {comp_status.overall_compliance:.1f}% "
                    f"({comp_status.requirements_met}/{comp_status.requirements_total} requirements met)",
                    styles['BodyText']
                ))
                if comp_status.security_level:
                    story.append(Paragraph(
                        f"Security Level: <b>{comp_status.security_level}</b>",
                        styles['BodyText']
                    ))
                story.append(Spacer(1, 0.1*inch))

        story.append(PageBreak())

        # Critical Devices
        if report.critical_devices:
            story.append(Paragraph("3. CRITICAL DEVICES", heading1_style))
            story.append(Spacer(1, 0.1*inch))

            for device in report.critical_devices[:10]:
                device_data = [
                    ['IP Address', device.ip],
                    ['Device Type', device.device_type],
                    ['Vendor/Model', f"{device.vendor} {device.model}"],
                    ['Risk Score', f"{device.risk_score:.1f}/100"],
                    ['Risk Level', device.risk_level.value],
                    ['CVE Count', str(device.cve_count)],
                ]

                device_table = Table(device_data, colWidths=[2*inch, 4*inch])
                device_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ecf0f1')),
                    ('FONT', (0, 0), (0, -1), 'Helvetica-Bold', 10),
                    ('FONT', (1, 0), (1, -1), 'Helvetica', 10),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]))
                story.append(device_table)
                story.append(Spacer(1, 0.15*inch))

        story.append(PageBreak())

        # Action Plan
        story.append(Paragraph("4. REMEDIATION ACTION PLAN", heading1_style))
        story.append(Spacer(1, 0.1*inch))

        if report.immediate_actions:
            story.append(Paragraph("4.1 IMMEDIATE Actions (Within 24 hours)", heading2_style))
            for i, action in enumerate(report.immediate_actions, 1):
                story.append(Paragraph(
                    f"<b>{i}. {action.title}</b>",
                    styles['BodyText']
                ))
                story.append(Paragraph(
                    f"   {action.description}",
                    styles['BodyText']
                ))
                if action.affected_devices:
                    story.append(Paragraph(
                        f"   Affected: {', '.join(action.affected_devices[:3])}",
                        styles['BodyText']
                    ))
                story.append(Spacer(1, 0.1*inch))

        if report.short_term_actions:
            story.append(Paragraph("4.2 SHORT TERM Actions (Within 1 week)", heading2_style))
            for i, action in enumerate(report.short_term_actions[:5], 1):
                story.append(Paragraph(
                    f"<b>{i}. {action.title}</b>",
                    styles['BodyText']
                ))
                story.append(Spacer(1, 0.05*inch))

        if report.medium_term_actions:
            story.append(Paragraph("4.3 MEDIUM TERM Actions (Within 1 month)", heading2_style))
            for i, action in enumerate(report.medium_term_actions[:5], 1):
                story.append(Paragraph(
                    f"{i}. {action.title}",
                    styles['BodyText']
                ))

        # Build PDF
        doc.build(story)

        return str(filepath)

    def generate_docx_report(
        self,
        report: RiskAssessmentReport,
        filename: Optional[str] = None
    ) -> str:
        """
        Generate DOCX (Word) pentest report

        Returns: Path to generated DOCX file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")

        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"risk_assessment_{timestamp}.docx"

        filepath = self.output_dir / filename

        # Create document
        doc = Document()

        # Set document properties
        doc.core_properties.title = "OT/ICS Security Penetration Test Report"
        doc.core_properties.author = "ICSScout"
        doc.core_properties.comments = f"Risk Assessment Report - {report.report_id}"

        # Title Page
        title = doc.add_heading('OT/ICS SECURITY', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading('PENETRATION TEST REPORT', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Risk Level
        risk_para = doc.add_paragraph()
        risk_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        risk_run = risk_para.add_run(f'Risk Level: {report.overall_risk_level.value}')
        risk_run.font.size = Pt(18)
        risk_run.font.bold = True
        risk_run.font.color.rgb = self._get_risk_color_rgb(report.overall_risk_level)

        doc.add_paragraph()

        # Report Information Table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'

        cells = table.rows[0].cells
        cells[0].text = 'Report ID:'
        cells[1].text = report.report_id

        cells = table.rows[1].cells
        cells[0].text = 'Assessment Date:'
        cells[1].text = report.scan_timestamp.strftime("%Y-%m-%d %H:%M:%S")

        cells = table.rows[2].cells
        cells[0].text = 'Total Devices Scanned:'
        cells[1].text = str(report.total_devices)

        cells = table.rows[3].cells
        cells[0].text = 'Overall Risk Score:'
        cells[1].text = f"{report.overall_risk_score:.2f}/100"

        # Bold first column
        for row in table.rows:
            row.cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_page_break()

        # Executive Summary
        doc.add_heading('1. EXECUTIVE SUMMARY', 1)

        exec_summary = f"""
        This report presents the results of a comprehensive security assessment of the OT/ICS network infrastructure.
        The assessment identified {report.total_devices} devices, with an overall risk rating of
        {report.overall_risk_level.value} (score: {report.overall_risk_score:.2f}/100).
        """
        doc.add_paragraph(exec_summary)

        # Key Findings
        doc.add_heading('Key Findings Summary', 2)

        findings_table = doc.add_table(rows=5, cols=3)
        findings_table.style = 'Light Grid Accent 1'

        # Header
        hdr_cells = findings_table.rows[0].cells
        hdr_cells[0].text = 'Severity'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Status'

        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True

        # Data
        row_cells = findings_table.rows[1].cells
        row_cells[0].text = 'Critical'
        row_cells[1].text = str(report.critical_count)
        row_cells[2].text = 'Immediate Action Required'

        row_cells = findings_table.rows[2].cells
        row_cells[0].text = 'High'
        row_cells[1].text = str(report.high_count)
        row_cells[2].text = 'Urgent'

        row_cells = findings_table.rows[3].cells
        row_cells[0].text = 'Medium'
        row_cells[1].text = str(len([f for f in report.all_findings if f.severity.value == "MEDIUM"]))
        row_cells[2].text = 'Important'

        row_cells = findings_table.rows[4].cells
        row_cells[0].text = 'Low'
        row_cells[1].text = str(len([f for f in report.all_findings if f.severity.value == "LOW"]))
        row_cells[2].text = 'Monitor'

        doc.add_page_break()

        # Detailed Assessment
        doc.add_heading('2. DETAILED ASSESSMENT', 1)

        # Network Security
        doc.add_heading('2.1 Network Security Assessment', 2)
        net_score = report.network_assessment.score
        net_risk = RiskLevel.from_score(100 - net_score)
        doc.add_paragraph(f"Score: {net_score:.1f}/100 | Risk Level: {net_risk.value}")

        if report.network_assessment.findings:
            doc.add_paragraph('Findings:', style='List Bullet')
            for finding in report.network_assessment.findings[:5]:
                doc.add_paragraph(f"{finding.title}: {finding.description}", style='List Bullet 2')

        # Device Security
        doc.add_heading('2.2 Device Security Assessment', 2)
        dev_score = report.device_assessment.score
        dev_risk = RiskLevel.from_score(100 - dev_score)
        doc.add_paragraph(f"Score: {dev_score:.1f}/100 | Risk Level: {dev_risk.value}")

        if report.device_assessment.findings:
            doc.add_paragraph('Findings:', style='List Bullet')
            for finding in report.device_assessment.findings[:5]:
                doc.add_paragraph(finding.title, style='List Bullet 2')

        # Vulnerability Assessment
        doc.add_heading('2.3 Vulnerability Assessment', 2)
        vuln_score = report.vulnerability_assessment.score
        vuln_risk = RiskLevel.from_score(100 - vuln_score)
        doc.add_paragraph(f"Score: {vuln_score:.1f}/100 | Risk Level: {vuln_risk.value}")

        if report.critical_findings:
            doc.add_paragraph('Critical Vulnerabilities:', style='List Bullet')
            for finding in report.critical_findings[:5]:
                cve_text = f" ({finding.cve_id})" if finding.cve_id else ""
                doc.add_paragraph(f"{finding.title}{cve_text}: {finding.description}", style='List Bullet 2')

        # Compliance Assessment
        doc.add_heading('2.4 Compliance Assessment', 2)
        comp_score = report.compliance_assessment.score
        doc.add_paragraph(f"Compliance Score: {comp_score:.1f}%")

        if report.compliance_status:
            for comp_status in report.compliance_status:
                doc.add_paragraph(
                    f"{comp_status.framework}: {comp_status.overall_compliance:.1f}% "
                    f"({comp_status.requirements_met}/{comp_status.requirements_total} requirements met)"
                )
                if comp_status.security_level:
                    doc.add_paragraph(f"Security Level: {comp_status.security_level}")

        doc.add_page_break()

        # Critical Devices
        if report.critical_devices:
            doc.add_heading('3. CRITICAL DEVICES', 1)

            for device in report.critical_devices[:10]:
                doc.add_heading(f"Device: {device.ip}", 3)

                device_table = doc.add_table(rows=6, cols=2)
                device_table.style = 'Light Grid'

                device_table.rows[0].cells[0].text = 'IP Address'
                device_table.rows[0].cells[1].text = device.ip

                device_table.rows[1].cells[0].text = 'Device Type'
                device_table.rows[1].cells[1].text = device.device_type

                device_table.rows[2].cells[0].text = 'Vendor/Model'
                device_table.rows[2].cells[1].text = f"{device.vendor} {device.model}"

                device_table.rows[3].cells[0].text = 'Risk Score'
                device_table.rows[3].cells[1].text = f"{device.risk_score:.1f}/100"

                device_table.rows[4].cells[0].text = 'Risk Level'
                device_table.rows[4].cells[1].text = device.risk_level.value

                device_table.rows[5].cells[0].text = 'CVE Count'
                device_table.rows[5].cells[1].text = str(device.cve_count)

                # Bold first column
                for row in device_table.rows:
                    row.cells[0].paragraphs[0].runs[0].font.bold = True

                doc.add_paragraph()

        doc.add_page_break()

        # Action Plan
        doc.add_heading('4. REMEDIATION ACTION PLAN', 1)

        if report.immediate_actions:
            doc.add_heading('4.1 IMMEDIATE Actions (Within 24 hours)', 2)
            for i, action in enumerate(report.immediate_actions, 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(action.title).bold = True
                doc.add_paragraph(action.description, style='List Bullet 2')
                if action.affected_devices:
                    doc.add_paragraph(f"Affected: {', '.join(action.affected_devices[:3])}", style='List Bullet 2')

        if report.short_term_actions:
            doc.add_heading('4.2 SHORT TERM Actions (Within 1 week)', 2)
            for i, action in enumerate(report.short_term_actions[:5], 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(action.title).bold = True

        if report.medium_term_actions:
            doc.add_heading('4.3 MEDIUM TERM Actions (Within 1 month)', 2)
            for i, action in enumerate(report.medium_term_actions[:5], 1):
                doc.add_paragraph(action.title, style='List Number')

        # Save document
        doc.save(str(filepath))

        return str(filepath)

    def generate_json_report(
        self,
        report: RiskAssessmentReport,
        filename: Optional[str] = None
    ) -> str:
        """Generate JSON report"""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"risk_assessment_{timestamp}.json"

        filepath = self.output_dir / filename

        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(report.to_dict(), f, indent=2, ensure_ascii=False)

        return str(filepath)

    def _get_risk_color(self, risk_level: RiskLevel) -> str:
        """Get hex color for risk level"""
        return risk_level.color

    def _get_risk_color_rgb(self, risk_level: RiskLevel) -> RGBColor:
        """Get RGB color for risk level (for DOCX)"""
        hex_color = risk_level.color.lstrip('#')
        r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
        return RGBColor(r, g, b)
